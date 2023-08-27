from pdfminer.high_level import extract_text
from docx.api import Document

def heading_from_pdf(pdf):
    #reader = PdfReader(pdf)
    response = extract_text(pdf)
    #for t in range(len(reader.pages)):
    #    text = reader.pages[t]
    #    response = response + text.extract_text()

    heading = "-1"
    s = 0
    for t in range(len(response)):
        if response[t] == '\n' and t < 73 and s > 0:
            heading = response[:t]
            response = response[t:]
            break
        if (response[t] >= 'a' and response[t] <= 'z') or (response[t] >= 'A' and response[t] <= 'Z'):
            s = 1
    return heading

def text_from_pdf(pdf):
    #reader = PdfReader(pdf)
    response = extract_text(pdf)
    #for t in range(len(reader.pages)):
    #    text = reader.pages[t]
    #    response = response + text.extract_text()

    s = 0
    for t in range(len(response)):
        if response[t] == '\n' and t < 73 and s > 0:
            response = response[t:]
            break
        if (response[t] >= 'a' and response[t] <= 'z') or (response[t] >= 'A' and response[t] <= 'Z'):
            s = 1
    return response

def heading_from_docx(docx):
    doc = Document(docx)
    texts = doc.paragraphs
    response = ""
    for text in texts:
        response = response + text.text
    heading = "-1"
    s = 0
    for t in range(len(response)):
        if response[t] == '\n' and t < 73 and s > 0:
            heading = response[:t]
            response = response[t:]
            break
        if (response[t] >= 'a' and response[t] <= 'z') or (response[t] >= 'A' and response[t] <= 'Z'):
            s = 1
    print(heading)
    return heading

def text_from_docx(docx):
    doc = Document(docx)
    texts = doc.paragraphs
    response = ""
    for text in texts:
        response = response + text.text
    s = 0
    for t in range(len(response)):
        if response[t] == '\n' and t < 73 and s > 0:
            response = response[t:]
            break
        if (response[t] >= 'a' and response[t] <= 'z') or (response[t] >= 'A' and response[t] <= 'Z'):
            s = 1
    return response

